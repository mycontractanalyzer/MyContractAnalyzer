import io

from docx import Document


def generate_lawyer_pack(report: str, contract_text: str, email: str,
                         title: str = "Договор") -> bytes:
    doc = Document()
    doc.add_heading("Пакет для юриста — MyContractAnalyzer", level=0)
    doc.add_paragraph(f"Клиент: {email} · Документ: {title}")
    doc.add_paragraph("Ниже — отчёт ИИ-аналитика и полный текст договора для профессиональной проверки.")
    doc.add_heading("Часть 1. Отчёт ИИ-аналитика", level=1)
    for line in report.splitlines():
        if line.strip():
            doc.add_paragraph(line.replace("**", ""))
    doc.add_heading("Часть 2. Полный текст договора", level=1)
    for line in (contract_text or "").splitlines():
        if line.strip():
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_protocol_docx(rows, email: str, title: str = "Договор") -> bytes:
    doc = Document()
    doc.add_heading("Протокол разногласий — MyContractAnalyzer", level=0)
    doc.add_paragraph(f"От: {email} · Документ: {title}")
    doc.add_paragraph("Предложения основаны на ИИ-анализе договора и открыты к обсуждению.")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Пункт / предмет"
    hdr[1].text = "Текущая редакция (риск)"
    hdr[2].text = "Предлагаемая редакция"
    for r in rows:
        cells = table.add_row().cells
        cells[0].text = str(r.get("clause", ""))
        cells[1].text = str(r.get("current", ""))
        cells[2].text = str(r.get("proposed", ""))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()