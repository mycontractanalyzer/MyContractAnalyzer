import io
import re

from docx import Document


def _add_rich(paragraph, text):
    parts = re.split(r"\*\*(.+?)\*\*", text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        if i % 2 == 1:
            run.bold = True


def _render_report(doc, report):
    for line in report.splitlines():
        line = line.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].replace("**", ""), level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:].replace("**", ""), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].replace("**", ""), level=1)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_rich(p, line[2:])
        else:
            p = doc.add_paragraph()
            _add_rich(p, line)


def generate_report_docx(report: str, email: str) -> bytes:
    doc = Document()
    doc.add_heading("MyContractAnalyzer — отчёт по договору", level=0)
    doc.add_paragraph(f"Запрошен пользователем: {email}")
    doc.add_paragraph("Не является юридической консультацией.")
    _render_report(doc, report)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_lawyer_pack(report: str, contract_text: str, email: str,
                         title: str = "Договор") -> bytes:
    doc = Document()
    doc.add_heading("Пакет для юриста — MyContractAnalyzer", level=0)
    doc.add_paragraph(f"Клиент: {email} · Документ: {title}")
    doc.add_paragraph("Ниже — отчёт ИИ-аналитика и полный текст договора для профессиональной проверки.")
    doc.add_heading("Часть 1. Отчёт ИИ-аналитика", level=1)
    _render_report(doc, report)
    doc.add_heading("Часть 2. Полный текст договора", level=1)
    for line in (contract_text or "").splitlines():
        if line.strip():
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()